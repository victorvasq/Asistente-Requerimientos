import const as const
import streamlit as st
import time
import asyncio
from langchain.chat_models import ChatOpenAI
from langchain.schema import SystemMessage
from langchain.prompts import ChatPromptTemplate, HumanMessagePromptTemplate, MessagesPlaceholder, PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.chains import LLMChain
from langchain.llms import OpenAI
from langchain.callbacks import get_openai_callback
from docx import Document
from io import BytesIO
from PIL import Image



# Función para crear un archivo Word
def create_word_file(titulo, contenido):
    doc = Document()
    doc.add_heading(titulo, 0)
    doc.add_paragraph(contenido)
    return doc

def formatoWord(contenido):
    contenido = contenido.replace("\n1.- ", "\n\n1.- ").replace("\n2.- ", "\n\n2.- ").replace("\n3.- ", "\n\n3.- ").replace("\n4.- ", "\n\n4.- ")
    return contenido

# Función para descargar el archivo Word
def download_word_file(doc):
    with BytesIO() as byte_io:
        doc.save(byte_io)
        byte_io.seek(0)
        return byte_io.read()

def modeloMemoryLangChainOpenAI(api_key, modelo, contextoSystem):
    prompt = ChatPromptTemplate.from_messages([
        MessagesPlaceholder(variable_name="chat_history"), # Donde se guardará la memoria.
        HumanMessagePromptTemplate.from_template("{human_input}"), 
        SystemMessage(content=contextoSystem), # Mensaje persistente del sistema
    ])
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    
    llm = ChatOpenAI(temperature=0.7, openai_api_key = api_key, model=modelo)
    chat_llm_chain = LLMChain(
        llm=llm,
        prompt=prompt,
        verbose=False,
        memory=memory,
    )
    return chat_llm_chain


def modeloOpenAI(api_key, modelo, contextoTemplate):
    template = contextoTemplate
    prompt_template = PromptTemplate(input_variables=["entrada"],
                                 template=template)
    llm = OpenAI(temperature=0.7, openai_api_key = api_key, model=modelo)
    chain = LLMChain(llm=llm, prompt=prompt_template)
    
    return chain


def imprimeChat():
    for msg in ss.messages:
        role = msg["role"]
        content = msg["content"].replace("\n", "<br>")
        if role == "assistant":
            st.markdown(f"**👩‍🦰 <span style='color: #246ba9;'>Asistente</span>**: {content}", unsafe_allow_html=True)
        else:
            st.markdown(f"**🙂 <span style='color: #008000;'>Usuario</span>**: {content}", unsafe_allow_html=True)

# Llamado del boton Chat
def form_callback_chat():
    ss["submit_btnChat"] = True

def form_callback_EditarResumen():
    ss["proceso"] = "EditaResumen"

def form_callback_ModificaResumen():
    ss.resumen = ss.user_input_Resumen
    ss["proceso"] = "Resumen"

def form_callback_CancelaModificaResumen():
    ss["proceso"] = "Resumen"

def form_callback_FinalizarProceso():
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    #st.experimental_rerun()

def form_callback_ContinuarProceso():
    ss["proceso"] = "Resumen"

def form_callback_DeseaFinalizar():
    ss["proceso"] = "Finalizar"



#####################################
# __main__


#Inicializar
ss = st.session_state
if "proceso" not in ss:
    ss["proceso"] = "Chat"
if "total_tokens" not in ss:
    ss["total_tokens"] = 0
if "submit_btnChat" not in ss:
    ss["submit_btnChat"] = False
if "resumen" not in ss:
    ss["resumen"] = ""
image = "loading.gif"

#openai_api_key = st.secrets["api_key"]


st.title("Gestión de Requerimientos") 


# Menu Lateral
with st.sidebar:
    openai_api_key = st.text_input("OpenAI API Key", key="chatbot_api_key", type="password")
    st.sidebar.header('Asistente de Requerimientos')
    st.markdown("""
        <div style="text-align: justify;">
            Bienvenido al asistente virtual de levantamiento de requerimientos de software. Este chatbot está diseñado para guiarlo eficientemente a través del proceso, capturando sus necesidades y generando como resultado, un documento con los objetivos y requisitos generales del proyecto. Este asistente ha sido creado para demostrar un caso de uso práctico, generando únicamente un documento básico como muestra de su funcionalidad.
        </div>
    """, unsafe_allow_html=True)
    st.text("")
    st.sidebar.header('Aplicación creada por \n')
    "[Víctor Vásquez](https://www.linkedin.com/in/victorvasquezrivas/)"
    "[victorvasquezrivas@gmail.com](mailto:victorvasquezrivas@gmail.com)"


### Chat ###
if ss["proceso"] == "Chat":
    st.markdown("Ingrese su requerimiento tecnológico, nuestra asistente lo guiará en su gestión.", unsafe_allow_html=True)
    
    if "chat_llm_chain" not in ss and openai_api_key:
        ss["chat_llm_chain"] = modeloMemoryLangChainOpenAI(openai_api_key, "gpt-4", const.CONTEXTO_SYSTEM_CHAT)
        
    # Guarda el primer mensaje del chat
    if "messages" not in ss:
        ss["messages"] = [{"role": "assistant", "content": "¿Qué requimiento de software necesita?"}]

    # Imprime el chat
    imprimeChat()

    # Click button chat
    if ss["submit_btnChat"] and ss.user_input_area != "":
        if not openai_api_key:
            st.info("Para continuar con el chat, debe ingresar la API Key.")
        else:
            
            respuestaUsuario = ss.user_input_area
            ss.user_input_area = ""
            respuest = respuestaUsuario.replace("\n", "<br>")
            ss.messages.append({"role": "user", "content": respuest})
            st.markdown(f"**🙂 <span style='color: #008000;'>Usuario</span>**: {respuest}", unsafe_allow_html=True)
            #st.markdown(f"**👩‍🦰 <span style='color: #246ba9;'>Asistente</span>**:", unsafe_allow_html=True)
            #st.image(image, width=50)

            message_placeholder = st.empty()
            message_placeholder.markdown("**👩‍🦰 <span style='color: #246ba9;'>Asistente</span>**: ... <img src='https://drive.google.com/uc?export=download&id=1ck9DDogy15NJxnNXKy8FiROZHwwnvVFS' alt='0' style='width: 25px;'>", unsafe_allow_html=True)
            
            # llamando al LLM
            chat_llm_chain = ss["chat_llm_chain"]
            with get_openai_callback() as cb: # para contar los tokens
                respuesta = chat_llm_chain.predict(human_input=respuestaUsuario)
            ss["total_tokens"] += cb.total_tokens
            
            if respuesta.startswith("Resumen:"):
                ss["proceso"] = "Resumen"
                respuesta = respuesta[8:] #Elimina la palabra Resumen:
                
                ss["resumen"] = formatoWord(respuesta.strip()) #contenido
                st.experimental_rerun()
            else:
                #st.image(image, width=1)
                ss.messages.append({"role": "assistant", "content": respuesta})
                
                
                respuesta = respuesta.replace("\n", "<br>")
                palabras = respuesta.split()
                result=""
                for palabra in palabras:
                    result += " " + palabra
                    message_placeholder.markdown("**👩‍🦰 <span style='color: #246ba9;'>Asistente</span>**: " + result + "▌", unsafe_allow_html=True)
                    time.sleep(0.05)
                message_placeholder.markdown("**👩‍🦰 <span style='color: #246ba9;'>Asistente</span>**: " + result, unsafe_allow_html=True)
                #st.experimental_rerun()
    
    
    # Formulario textArea
    with st.form(key='my_form'):
        text = st.text_area("Usuario:", '',key="user_input_area")
        st.form_submit_button(label='Enviar', on_click=form_callback_chat)



### Finalizar ###
if ss["proceso"] == "Finalizar":
    with st.container():
        st.info("Al finalizar se eliminará toda la información ingresada, procure descargar el word.  ¿Desea Finalizar el proceso?")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.button("Seguir Modificando", on_click=form_callback_ContinuarProceso)
        with col2:
            st.download_button(
                label="Descargar Word",
                data=BytesIO(download_word_file(create_word_file("Gestión de Requerimientos", ss["resumen"]))),
                file_name="requerimiento.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col3:
            st.button("Finalizar", on_click=form_callback_FinalizarProceso, type="primary")



### Resumen ###
if ss["proceso"] == "Resumen":
    with st.container():
        st.info("Revise la información y verifique si necesita modificarla")
        if "resumen" in ss:
            content = ss["resumen"].replace("\n", "<br>").replace(" ", "&nbsp;")
            st.markdown(f"{content}", unsafe_allow_html=True)

        st.markdown("<hr/>", unsafe_allow_html=True)
        
        left_column, center_column, rigth_column = st.columns(3)
        with left_column:
            st.button('Modificar Información', on_click=form_callback_EditarResumen, type="secondary")
        with center_column:
            st.download_button(
                label="Descargar Word",
                data=BytesIO(download_word_file(create_word_file("Requerimiento Tecnológico", ss["resumen"]))),
                file_name="requerimiento.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with rigth_column:
            st.button('Finalizar', on_click=form_callback_DeseaFinalizar, type="primary")
            
            


### Edita Resumen ###
if ss["proceso"] == "EditaResumen":
    content = ss["resumen"]
    with st.container():
        text = st.text_area('Modifica la información aquí:', content,key="user_input_Resumen", height=500)
        
        left_column, rigth_column = st.columns(2)
        with left_column:
            st.button('Guardar', on_click=form_callback_ModificaResumen, type="primary")
        with rigth_column:
            st.button('Cancelar', on_click=form_callback_CancelaModificaResumen, type="secondary")


st.markdown("<hr/>", unsafe_allow_html=True)
st.write("Total Tokens: ", ss["total_tokens"])
